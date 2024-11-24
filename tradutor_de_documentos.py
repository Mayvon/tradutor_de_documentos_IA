# -*- coding: utf-8 -*-
"""tradutor_de_documentos_docx.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/gist/Mayvon/613b210a3c7989a99b36f79dcccfdd66/tradutor_de_documentos_docx.ipynb
"""

!pip install requests python-docx

import requests
from docx import Document
import os

subscription_key = '3kFOfGE1cu2pxHopjdYGj7sNaGZLMTmgyY7v4vOzp0uolCH2QGnJJQQJ99AKAC8vTInXJ3w3AAAbACOGLVv8'
endpoit = 'https://api.cognitive.microsofttranslator.com/'
location = 'westus2'
language_destination = 'pt-br'

def translator_text(text, language_destination):
 path = '/translate'
 constructed_languge_url = endpoit + path
 headers = {
     'Ocp-Apim-Subscription-Key': subscription_key,
     'Ocp-Apim-Subscription-Region': location,
     'Content-type': 'application/json',
     'X-ClientTraceId': str(os.urandom(16))
 }


 body = [{
     'text': text
 }]
 params = {
     'api-version': '3.0',
     'from': 'en',
     'to': [language_destination]
 }

 request = requests.post(constructed_languge_url, headers=headers, json=body, params=params)
 response = request.json()
 return response[0]["translations"][0]["text"]

def translated_document (path):
    document = Document(path)
    full_text = []
    for paragraph in document.paragraphs:

      translated_paragraph = translator_text(paragraph.text, language_destination)
      full_text.append(translated_paragraph)


    translated_doc = Document()
    for line in full_text:
      translated_doc.add_paragraph(line)
    path_translated = path.replace('.docx', f"_{language_destination}.docx")
    translated_doc.save(path_translated)
    return path_translated

input_file = "/content/arquivo.docx"
    translated_document(input_file)